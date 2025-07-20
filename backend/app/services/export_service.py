"""Export service for generating PDF, DOCX and other formats."""

import io
import logging
import os
from datetime import datetime
from typing import Any

import markdown  # type: ignore[import-untyped]
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors  # type: ignore[import-untyped]
from reportlab.lib.pagesizes import A4  # type: ignore[import-untyped]
from reportlab.lib.styles import (  # type: ignore[import-untyped]
    ParagraphStyle,
    getSampleStyleSheet,
)
from reportlab.lib.units import inch  # type: ignore[import-untyped]
from reportlab.pdfbase import pdfmetrics  # type: ignore[import-untyped]
from reportlab.pdfbase.ttfonts import TTFont  # type: ignore[import-untyped]
from reportlab.platypus import (  # type: ignore[import-untyped]
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

logger = logging.getLogger(__name__)


class ExportService:
    """导出服务类."""

    def __init__(self) -> None:
        """初始化导出服务."""
        self.styles = getSampleStyleSheet()
        self._setup_pdf_fonts()
        self._setup_custom_styles()

    def _setup_pdf_fonts(self) -> None:
        """设置PDF中文字体."""
        try:
            # 尝试注册中文字体
            font_paths = [
                "/System/Library/Fonts/STHeiti Light.ttc",  # macOS
                "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",  # Linux
                "C:/Windows/Fonts/simhei.ttf",  # Windows
            ]

            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont("ChineseFont", font_path))
                        self.chinese_font = "ChineseFont"
                        return
                    except Exception:
                        continue

            # 如果没有找到中文字体，使用默认字体
            self.chinese_font = "Helvetica"

        except Exception as e:
            logger.warning(f"Failed to setup Chinese font: {e}")
            self.chinese_font = "Helvetica"

    def _setup_custom_styles(self) -> None:
        """设置自定义样式."""
        # 标题样式
        self.styles.add(
            ParagraphStyle(
                name="ChineseTitle",
                parent=self.styles["Title"],
                fontName=self.chinese_font,
                fontSize=24,
                textColor=colors.HexColor("#1a1a1a"),
                spaceAfter=30,
                alignment=1,  # CENTER alignment
            )
        )

        # 正文样式
        self.styles.add(
            ParagraphStyle(
                name="ChineseBody",
                parent=self.styles["BodyText"],
                fontName=self.chinese_font,
                fontSize=11,
                leading=16,
                textColor=colors.HexColor("#333333"),
                spaceAfter=12,
            )
        )

        # 标题样式
        for i in range(1, 4):
            self.styles.add(
                ParagraphStyle(
                    name=f"ChineseHeading{i}",
                    parent=self.styles[f"Heading{i}"],
                    fontName=self.chinese_font,
                    fontSize=18 - i * 2,
                    textColor=colors.HexColor("#2c3e50"),
                    spaceAfter=12,
                    spaceBefore=12,
                )
            )

    async def export_note_to_pdf(
        self,
        note: dict[str, Any],
        include_metadata: bool = True,
        include_versions: bool = False,
        versions: list[dict[str, Any]] | None = None,
    ) -> bytes:
        """导出笔记为PDF格式."""
        buffer = io.BytesIO()

        # 创建PDF文档
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        # 构建内容
        story = []

        # 标题
        story.append(Paragraph(note["title"], self.styles["ChineseTitle"]))
        story.append(Spacer(1, 0.2 * inch))

        # 元数据
        if include_metadata:
            metadata_text = f"""
            <para>
            <b>创建时间:</b> {note["created_at"]}<br/>
            <b>更新时间:</b> {note["updated_at"]}<br/>
            <b>版本:</b> {note.get("version", 1)}<br/>
            <b>标签:</b> {", ".join(note.get("tags", []))}<br/>
            </para>
            """
            story.append(Paragraph(metadata_text, self.styles["ChineseBody"]))
            story.append(Spacer(1, 0.3 * inch))

        # 主要内容
        content_html = self._markdown_to_html(note["content"])
        content_paragraphs = self._html_to_reportlab(content_html)
        story.extend(content_paragraphs)

        # 版本历史
        if include_versions and versions:
            story.append(PageBreak())
            story.append(Paragraph("版本历史", self.styles["ChineseHeading1"]))
            story.append(Spacer(1, 0.2 * inch))

            for version in versions[:10]:  # 最多显示10个版本
                version_text = f"""
                <para>
                <b>版本 {version["version_number"]}</b> - {version["created_at"]}<br/>
                {version.get("change_summary", "无描述")}<br/>
                </para>
                """
                story.append(Paragraph(version_text, self.styles["ChineseBody"]))
                story.append(Spacer(1, 0.1 * inch))

        # 生成PDF
        doc.build(story)

        buffer.seek(0)
        return buffer.getvalue()

    async def export_note_to_docx(
        self,
        note: dict[str, Any],
        include_metadata: bool = True,
        include_versions: bool = False,
        versions: list[dict[str, Any]] | None = None,
    ) -> bytes:
        """导出笔记为DOCX格式."""
        # 创建文档
        doc = DocxDocument()

        # 设置文档标题
        doc.core_properties.title = note["title"]
        doc.core_properties.author = note.get("author", "Second Brain")
        doc.core_properties.created = datetime.fromisoformat(note["created_at"])

        # 添加标题
        title = doc.add_heading(note["title"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加元数据
        if include_metadata:
            doc.add_paragraph()
            metadata = doc.add_paragraph()
            metadata.add_run("创建时间: ").bold = True
            metadata.add_run(note["created_at"])
            metadata.add_run("\n")
            metadata.add_run("更新时间: ").bold = True
            metadata.add_run(note["updated_at"])
            metadata.add_run("\n")
            metadata.add_run("版本: ").bold = True
            metadata.add_run(str(note.get("version", 1)))
            metadata.add_run("\n")
            metadata.add_run("标签: ").bold = True
            metadata.add_run(", ".join(note.get("tags", [])))
            doc.add_paragraph()

        # 添加主要内容
        self._add_markdown_to_docx(doc, note["content"])

        # 添加版本历史
        if include_versions and versions:
            doc.add_page_break()
            doc.add_heading("版本历史", 1)

            # 创建版本表格
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"

            # 表头
            header_cells = table.rows[0].cells
            header_cells[0].text = "版本"
            header_cells[1].text = "时间"
            header_cells[2].text = "说明"

            # 添加版本数据
            for version in versions[:10]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(version["version_number"])
                row_cells[1].text = version["created_at"]
                row_cells[2].text = version.get("change_summary", "无描述")

        # 保存到内存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    async def export_document_to_pdf(
        self,
        document: dict[str, Any],
        include_annotations: bool = True,
        annotations: list[dict[str, Any]] | None = None,
    ) -> bytes:
        """导出文档为PDF格式."""
        buffer = io.BytesIO()

        # 创建PDF文档
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        story = []

        # 标题
        story.append(
            Paragraph(
                document.get("title", document["filename"]), self.styles["ChineseTitle"]
            )
        )
        story.append(Spacer(1, 0.2 * inch))

        # 文档信息
        info_text = f"""
        <para>
        <b>文件名:</b> {document["filename"]}<br/>
        <b>上传时间:</b> {document["created_at"]}<br/>
        <b>文件大小:</b> {self._format_file_size(document["file_size"])}<br/>
        </para>
        """
        story.append(Paragraph(info_text, self.styles["ChineseBody"]))
        story.append(Spacer(1, 0.3 * inch))

        # 文档内容
        if document.get("content"):
            content_paragraphs = self._process_document_content(document["content"])
            story.extend(content_paragraphs)

        # 标注
        if include_annotations and annotations:
            story.append(PageBreak())
            story.append(Paragraph("标注列表", self.styles["ChineseHeading1"]))
            story.append(Spacer(1, 0.2 * inch))

            for ann in annotations:
                ann_text = f"""
                <para>
                <b>{ann["type"]}</b> - 第{ann.get("page_number", "?")}页<br/>
                {ann.get("text", "")}
                {f"<br/><i>备注: {ann['comment']}</i>" if ann.get("comment") else ""}
                </para>
                """
                story.append(Paragraph(ann_text, self.styles["ChineseBody"]))
                story.append(Spacer(1, 0.1 * inch))

        # 生成PDF
        doc.build(story)

        buffer.seek(0)
        return buffer.getvalue()

    async def export_space_to_pdf(
        self,
        space: dict[str, Any],
        documents: list[dict[str, Any]],
        notes: list[dict[str, Any]],
        include_content: bool = False,
    ) -> bytes:
        """导出空间为PDF格式."""
        buffer = io.BytesIO()

        # 创建PDF文档
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        story = []

        # 空间标题
        story.append(Paragraph(space["name"], self.styles["ChineseTitle"]))
        story.append(Spacer(1, 0.2 * inch))

        # 空间描述
        if space.get("description"):
            story.append(Paragraph(space["description"], self.styles["ChineseBody"]))
            story.append(Spacer(1, 0.3 * inch))

        # 统计信息
        stats_text = f"""
        <para>
        <b>文档数量:</b> {len(documents)}<br/>
        <b>笔记数量:</b> {len(notes)}<br/>
        <b>创建时间:</b> {space["created_at"]}<br/>
        </para>
        """
        story.append(Paragraph(stats_text, self.styles["ChineseBody"]))
        story.append(Spacer(1, 0.5 * inch))

        # 文档列表
        if documents:
            story.append(Paragraph("文档列表", self.styles["ChineseHeading1"]))
            story.append(Spacer(1, 0.2 * inch))

            for document in documents:
                doc_text = f"• {document.get('title', document['filename'])}"
                story.append(Paragraph(doc_text, self.styles["ChineseBody"]))

                if include_content and document.get("summary"):
                    summary_text = f"  <i>{document['summary'][:200]}...</i>"
                    story.append(Paragraph(summary_text, self.styles["ChineseBody"]))

                story.append(Spacer(1, 0.1 * inch))

        # 笔记列表
        if notes:
            story.append(PageBreak())
            story.append(Paragraph("笔记列表", self.styles["ChineseHeading1"]))
            story.append(Spacer(1, 0.2 * inch))

            for note in notes:
                note_text = f"• {note['title']}"
                story.append(Paragraph(note_text, self.styles["ChineseBody"]))

                if include_content and note.get("content"):
                    content_preview = note["content"][:200] + "..."
                    preview_text = f"  <i>{content_preview}</i>"
                    story.append(Paragraph(preview_text, self.styles["ChineseBody"]))

                story.append(Spacer(1, 0.1 * inch))

        # 生成PDF
        doc.build(story)

        buffer.seek(0)
        return buffer.getvalue()

    def _markdown_to_html(self, markdown_text: str) -> str:
        """转换Markdown为HTML."""
        return markdown.markdown(
            markdown_text,
            extensions=["fenced_code", "tables", "footnotes", "nl2br", "codehilite"],
        )

    def _html_to_reportlab(self, html_content: str) -> list[Any]:
        """转换HTML为ReportLab元素."""
        # 简化版HTML解析
        paragraphs = []

        # 分段处理
        lines = html_content.split("\n")
        for line in lines:
            if line.strip():
                # 简单处理，实际应使用HTML解析器
                clean_line = line.replace("<p>", "").replace("</p>", "")
                clean_line = clean_line.replace("<br>", "<br/>")

                try:
                    para = Paragraph(clean_line, self.styles["ChineseBody"])
                    paragraphs.append(para)
                    paragraphs.append(Spacer(1, 0.1 * inch))
                except Exception:
                    # 如果解析失败，使用纯文本
                    para = Paragraph(self._escape_xml(line), self.styles["ChineseBody"])
                    paragraphs.append(para)
                    paragraphs.append(Spacer(1, 0.1 * inch))

        return paragraphs

    def _add_markdown_to_docx(self, doc: Any, markdown_text: str) -> None:
        """添加Markdown内容到DOCX文档."""
        # 简化版Markdown处理
        lines = markdown_text.split("\n")

        for line in lines:
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                # 列表项
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("1. ") or (
                len(line) >= 3 and line[0].isdigit() and line[1:3] == ". "
            ):
                # 编号列表
                doc.add_paragraph(line[3:], style="List Number")
            elif line.strip():
                # 普通段落
                doc.add_paragraph(line)
            else:
                # 空行
                doc.add_paragraph()

    def _process_document_content(self, content: str) -> list[Any]:
        """处理文档内容."""
        paragraphs = []

        # 按段落分割
        for para in content.split("\n\n"):
            if para.strip():
                try:
                    p = Paragraph(self._escape_xml(para), self.styles["ChineseBody"])
                    paragraphs.append(p)
                    paragraphs.append(Spacer(1, 0.2 * inch))
                except Exception:
                    continue

        return paragraphs

    def _escape_xml(self, text: str) -> str:
        """转义XML特殊字符."""
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&apos;")
        )

    def _format_file_size(self, size_bytes: int) -> str:
        """格式化文件大小."""
        size: float = float(size_bytes)
        for unit in ["B", "KB", "MB", "GB"]:
            if size < 1024.0:
                return f"{size:.1f} {unit}"
            size /= 1024.0
        return f"{size:.1f} TB"

    async def export_notes_to_pdf(
        self,
        notes: list[dict[str, Any]],
        title: str = "笔记合集",
        include_metadata: bool = True,
    ) -> bytes:
        """导出多个笔记为单个PDF文件."""
        buffer = io.BytesIO()

        # 创建PDF文档
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        story = []

        # 合集标题
        story.append(Paragraph(title, self.styles["ChineseTitle"]))
        story.append(Spacer(1, 0.3 * inch))

        # 目录
        story.append(Paragraph("目录", self.styles["ChineseHeading1"]))
        for i, note in enumerate(notes, 1):
            toc_text = f"{i}. {note['title']}"
            story.append(Paragraph(toc_text, self.styles["ChineseBody"]))
        story.append(PageBreak())

        # 每个笔记内容
        for i, note in enumerate(notes, 1):
            # 笔记标题
            story.append(
                Paragraph(f"{i}. {note['title']}", self.styles["ChineseHeading1"])
            )
            story.append(Spacer(1, 0.2 * inch))

            # 元数据
            if include_metadata:
                meta_text = f"""
                <para>
                <b>创建时间:</b> {note["created_at"]}<br/>
                <b>更新时间:</b> {note["updated_at"]}<br/>
                <b>版本:</b> {note.get("version", 1)}<br/>
                <b>标签:</b> {", ".join(note.get("tags", []))}<br/>
                </para>
                """
                story.append(Paragraph(meta_text, self.styles["ChineseBody"]))
                story.append(Spacer(1, 0.3 * inch))

            # 笔记内容
            html_content = markdown.markdown(note["content"])
            content_elements = self._html_to_reportlab(html_content)
            story.extend(content_elements)

            # 分页（除了最后一个）
            if i < len(notes):
                story.append(PageBreak())

        # 生成PDF
        doc.build(story)

        buffer.seek(0)
        return buffer.getvalue()

    async def export_notes_to_docx(
        self,
        notes: list[dict[str, Any]],
        title: str = "笔记合集",
        include_metadata: bool = True,
    ) -> bytes:
        """导出多个笔记为单个DOCX文件."""
        # 创建DOCX文档
        doc = DocxDocument()

        # 合集标题
        doc.add_heading(title, 0)

        # 目录
        doc.add_heading("目录", 1)
        for i, note in enumerate(notes, 1):
            doc.add_paragraph(f"{i}. {note['title']}")

        # 分页
        doc.add_page_break()

        # 每个笔记内容
        for i, note in enumerate(notes, 1):
            # 笔记标题
            doc.add_heading(f"{i}. {note['title']}", 1)

            # 元数据
            if include_metadata:
                metadata = doc.add_paragraph()
                metadata.add_run("创建时间: ").bold = True
                metadata.add_run(f"{note['created_at']}\n")
                metadata.add_run("更新时间: ").bold = True
                metadata.add_run(f"{note['updated_at']}\n")
                metadata.add_run("版本: ").bold = True
                metadata.add_run(f"{note.get('version', 1)}\n")
                metadata.add_run("标签: ").bold = True
                metadata.add_run(", ".join(note.get("tags", [])))

                doc.add_paragraph()  # 空行

            # 笔记内容
            self._add_markdown_to_docx(doc, note["content"])

            # 分页（除了最后一个）
            if i < len(notes):
                doc.add_page_break()

        # 保存到内存
        buffer = io.BytesIO()
        doc.save(buffer)

        buffer.seek(0)
        return buffer.getvalue()

    async def export_documents_to_pdf(
        self,
        documents: list[dict[str, Any]],
        title: str = "文档合集",
        include_content: bool = True,
    ) -> bytes:
        """导出多个文档为单个PDF文件."""
        buffer = io.BytesIO()

        # 创建PDF文档
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        story = []

        # 合集标题
        story.append(Paragraph(title, self.styles["ChineseTitle"]))
        story.append(Spacer(1, 0.3 * inch))

        # 文档列表
        story.append(Paragraph("文档列表", self.styles["ChineseHeading1"]))
        for i, document in enumerate(documents, 1):
            list_text = f"{i}. {document['filename']} ({self._format_file_size(document.get('file_size', 0))})"
            story.append(Paragraph(list_text, self.styles["ChineseBody"]))
        story.append(PageBreak())

        # 每个文档内容
        for i, document in enumerate(documents, 1):
            # 文档标题
            story.append(
                Paragraph(
                    f"{i}. {document['filename']}", self.styles["ChineseHeading1"]
                )
            )
            story.append(Spacer(1, 0.2 * inch))

            # 文档信息
            info_text = f"""
            <para>
            <b>文件类型:</b> {document.get("mime_type", "unknown")}<br/>
            <b>文件大小:</b> {self._format_file_size(document.get("file_size", 0))}<br/>
            <b>上传时间:</b> {document["created_at"]}<br/>
            <b>标签:</b> {", ".join(document.get("tags", []))}<br/>
            </para>
            """
            story.append(Paragraph(info_text, self.styles["ChineseBody"]))
            story.append(Spacer(1, 0.3 * inch))

            # 文档内容
            if include_content and document.get("extracted_text"):
                story.append(Paragraph("内容摘要", self.styles["ChineseHeading2"]))
                content_elements = self._process_document_content(
                    document["extracted_text"][:1000] + "..."  # 限制内容长度
                )
                story.extend(content_elements)

            # 分页（除了最后一个）
            if i < len(documents):
                story.append(PageBreak())

        # 生成PDF
        doc.build(story)

        buffer.seek(0)
        return buffer.getvalue()

    async def export_conversation_to_json(
        self,
        conversation: dict[str, Any],
        messages: list[dict[str, Any]],
        include_branches: bool = False,
    ) -> dict[str, Any]:
        """导出对话为JSON格式."""
        export_data = {
            "conversation": {
                "id": conversation["id"],
                "title": conversation["title"],
                "space_id": conversation.get("space_id"),
                "created_at": conversation["created_at"],
                "updated_at": conversation["updated_at"],
                "agent_id": conversation.get("agent_id"),
                "model": conversation.get("model"),
                "search_enabled": conversation.get("search_enabled", False),
            },
            "messages": [],
            "statistics": {
                "total_messages": len(messages),
                "user_messages": sum(1 for m in messages if m["role"] == "user"),
                "assistant_messages": sum(
                    1 for m in messages if m["role"] == "assistant"
                ),
            },
        }

        # 处理消息
        for msg in messages:
            message_data = {
                "id": msg["id"],
                "role": msg["role"],
                "content": msg["content"],
                "created_at": msg["created_at"],
                "model": msg.get("model"),
                "parent_message_id": msg.get("parent_message_id"),
            }

            if include_branches and msg.get("branch_id"):
                message_data["branch_id"] = msg["branch_id"]

            export_data["messages"].append(message_data)

        return export_data


# 创建全局实例
export_service = ExportService()
